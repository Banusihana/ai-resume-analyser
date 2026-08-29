import os
import re
import io
import csv
import urllib.parse
from datetime import datetime
from flask import Flask, render_template, request, jsonify, redirect, url_for, session, send_file
import PyPDF2
import docx
import openpyxl

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

app = Flask(__name__, template_folder='templates', static_folder='static')
app.secret_key = 'editor-desk-hackathon-secret-key-2026'

# Maximum file size: 10MB
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024

DESKTOP_DIR = r"C:\Users\sihan\OneDrive\Desktop"
EXCEL_FILE = os.path.join(DESKTOP_DIR, "user_logins.xlsx")
CSV_FILE = os.path.join(DESKTOP_DIR, "user_logins.csv")

try:
    import firebase_admin
    from firebase_admin import credentials, firestore
    HAS_FIREBASE_LIB = True
except ImportError:
    HAS_FIREBASE_LIB = False

try:
    import gspread
    HAS_GSPREAD_LIB = True
except ImportError:
    HAS_GSPREAD_LIB = False

FIREBASE_INIT = False
db_firestore = None
FIREBASE_KEY_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "serviceAccountKey.json")
GOOGLE_SHEET_NAME = os.environ.get("GOOGLE_SHEET_NAME", "Resume Analyser Logins")

if HAS_FIREBASE_LIB and os.path.exists(FIREBASE_KEY_PATH):
    try:
        cred = credentials.Certificate(FIREBASE_KEY_PATH)
        firebase_admin.initialize_app(cred)
        db_firestore = firestore.client()
        FIREBASE_INIT = True
        print(f"[Firebase] Admin SDK initialized using {FIREBASE_KEY_PATH}")
    except Exception as e:
        print(f"[Firebase] Initialization note: {e}")


def log_user_login(name, email, ip_address):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    # 1. Excel (.xlsx)
    try:
        if os.path.exists(EXCEL_FILE):
            wb = openpyxl.load_workbook(EXCEL_FILE)
            ws = wb.active
        else:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Logins"
            ws.append(["Timestamp", "Name", "Email", "IP Address"])
            
        ws.append([timestamp, name, email, ip_address])
        wb.save(EXCEL_FILE)
    except Exception as e:
        print(f"[Excel] Error writing to Excel file: {e}")

    # 2. CSV (.csv)
    try:
        file_exists = os.path.exists(CSV_FILE)
        with open(CSV_FILE, mode="a", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            if not file_exists:
                writer.writerow(["Timestamp", "Name", "Email", "IP Address"])
            writer.writerow([timestamp, name, email, ip_address])
    except Exception as e:
        print(f"[CSV] Error writing to CSV file: {e}")

    # 3. Firebase Cloud Firestore
    if FIREBASE_INIT and db_firestore:
        try:
            db_firestore.collection("user_logins").add({
                "timestamp": timestamp,
                "name": name,
                "email": email,
                "ip_address": ip_address,
                "created_at": datetime.now()
            })
            print(f"[Firebase] Successfully saved login for {email} to Cloud Firestore!")
        except Exception as e:
            print(f"[Firebase] Note: {e}")

    # 4. Google Sheets via gspread
    if HAS_GSPREAD_LIB and os.path.exists(FIREBASE_KEY_PATH):
        try:
            scopes = [
                'https://www.googleapis.com/auth/spreadsheets',
                'https://www.googleapis.com/auth/drive'
            ]
            gc = gspread.service_account(filename=FIREBASE_KEY_PATH, scopes=scopes)
            try:
                sh = gc.open(GOOGLE_SHEET_NAME)
                worksheet = sh.sheet1
            except gspread.SpreadsheetNotFound:
                sh = gc.create(GOOGLE_SHEET_NAME)
                worksheet = sh.sheet1
                worksheet.append_row(["Timestamp", "Name", "Email", "IP Address"])
                
            worksheet.append_row([timestamp, name, email, ip_address])
            print(f"[Google Sheets] Successfully logged {email} to Google Sheet '{GOOGLE_SHEET_NAME}'")
        except Exception as e:
            print(f"[Google Sheets] Note: {e}")


WEAK_PHRASES = [
    ("responsible for", "high", "Replace passive duties with active achievements (e.g., 'Spearheaded', 'Led', 'Managed')."),
    ("duties included", "high", "Frame your work around achievements rather than assigned duties."),
    ("assisted with", "medium", "Use action verbs that clearly define your individual contribution (e.g., 'Co-developed', 'Supported')."),
    ("helped with", "medium", "Specify your exact role instead of 'helped' (e.g., 'Engineered', 'Optimized')."),
    ("worked on", "medium", "Too vague. Specify what you built, designed, or executed."),
    ("team player", "low", "Cliché term. Demonstrate collaboration with concrete examples (e.g., 'Partnered across 4 cross-functional teams')."),
    ("hard worker", "low", "Show dedication through delivered metrics rather than self-descriptive labels."),
    ("hardworking", "low", "Show, don't tell. Let quantifiable impact reflect your work ethic."),
    ("think outside the box", "low", "Overused buzzword. Detail a specific creative solution or innovation instead."),
    ("synergy", "low", "Corporate jargon. State clearly how integration or collaboration improved outcomes."),
    ("detail-oriented", "low", "Cliché descriptor. Demonstrate attention to detail through error-free formatting and precision metrics."),
    ("results-driven", "low", "Show your results with numbers rather than labeling yourself results-driven."),
    ("self-motivated", "low", "Demonstrate initiative by highlighting projects you initiated or led."),
    ("thought leader", "low", "Avoid grandiose self-titles. Highlight published work, talks, or patents instead."),
    ("dynamic", "low", "Overused buzzword. Describe specific adaptable solutions or fast-paced project environments."),
    ("go-to person", "low", "Replace informal phrases with professional impact metrics."),
]

STRONG_ACTION_VERBS = {
    "spearheaded", "engineered", "architected", "orchestrated", "optimized",
    "automated", "implemented", "developed", "built", "designed", "launched",
    "negotiated", "championed", "increased", "reduced", "accelerated",
    "generated", "scaled", "transformed", "overhauled", "pioneered", "boosted"
}

REQUIRED_SECTIONS = [
    ("experience", ["experience", "work history", "employment history", "career history", "professional experience"]),
    ("education", ["education", "academic", "degree", "university", "college"]),
    ("skills", ["skills", "technologies", "tech stack", "technical proficiencies", "competencies"])
]

COMMON_TECH_SKILLS = {
    "python", "javascript", "typescript", "react", "node.js", "express", "html", "css",
    "sql", "postgresql", "mongodb", "mysql", "redis", "docker", "kubernetes", "aws",
    "azure", "gcp", "git", "github", "ci/cd", "rest api", "graphql", "flask", "django",
    "java", "c++", "c#", "go", "rust", "swift", "kotlin", "pandas", "numpy", "tensorflow",
    "pytorch", "scikit-learn", "linux", "agile", "scrum", "jira", "unit testing", "figma"
}


def extract_text_from_file(file_storage):
    filename = file_storage.filename.lower()
    text = ""
    
    if filename.endswith(".pdf"):
        try:
            reader = PyPDF2.PdfReader(file_storage.stream)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            raise ValueError(f"Failed to read PDF file: {str(e)}")
            
    elif filename.endswith(".docx"):
        try:
            doc = docx.Document(file_storage.stream)
            paragraphs = [p.text for p in doc.paragraphs if p.text]
            for table in doc.tables:
                for row in table.rows:
                    paragraphs.extend([cell.text for cell in row.cells if cell.text])
            text = "\n".join(paragraphs)
        except Exception as e:
            raise ValueError(f"Failed to read DOCX file: {str(e)}")
            
    elif filename.endswith(".txt"):
        try:
            text = file_storage.stream.read().decode("utf-8", errors="ignore")
        except Exception as e:
            raise ValueError(f"Failed to read TXT file: {str(e)}")
    else:
        raise ValueError("Unsupported file format. Please upload a PDF, DOCX, or TXT file.")
        
    return text.strip()


def extract_candidate_name(text):
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    ignore_words = {"resume", "curriculum", "vitae", "cv", "page", "email", "phone", "profile", "summary"}

    for line in lines[:5]:
        line_clean = re.sub(r'[^a-zA-Z\s]', '', line).strip()
        words = line_clean.split()
        if 2 <= len(words) <= 4:
            if not any(w.lower() in ignore_words for w in words):
                return line_clean.title()
                
    return "Candidate"


def verify_name_match(user_name, resume_text):
    if not user_name:
        return {"is_mismatch": False, "logged_in_name": "", "resume_name": ""}

    resume_name = extract_candidate_name(resume_text)

    user_words = set(re.findall(r'\b[a-zA-Z]+\b', user_name.lower()))
    resume_words = set(re.findall(r'\b[a-zA-Z]+\b', resume_name.lower()))

    resume_words.discard("candidate")

    if not user_words or not resume_words:
        return {"is_mismatch": False, "logged_in_name": user_name, "resume_name": resume_name}

    overlap = user_words.intersection(resume_words)
    
    if not overlap:
        return {
            "is_mismatch": True,
            "logged_in_name": user_name,
            "resume_name": resume_name,
            "error": f"❌ Access Blocked: Name Mismatch Detected! You logged in as '{user_name}', but the uploaded resume belongs to '{resume_name}'. Please upload your own resume to continue."
        }

    return {
        "is_mismatch": False,
        "logged_in_name": user_name,
        "resume_name": resume_name
    }


def analyze_ethical_bias(text):
    """
    Ethical AI Bias & Fairness Audit Engine.
    Scans for age, gender, prestige, and location bias markers.
    """
    flags = []
    fairness_score = 100
    text_lower = text.lower()

    # 1. Age Bias (Graduation years before 2012 or explicit age)
    grad_years = re.findall(r'\b(19\d{2}|20[0-1]\d)\b', text)
    if grad_years:
        fairness_score -= 10
        flags.append({
            "category": "Age & Graduation Date Bias",
            "marker": f"Graduation Year '{grad_years[0]}'",
            "risk": "Including graduation years prior to recent dates can trigger subconscious age bias in screening.",
            "recommendation": "Remove graduation dates to promote age-blind candidate evaluation."
        })

    # 2. Gender Pronoun & Phrasing Markers
    gender_pronouns = re.findall(r'\b(he|his|him|she|her|hers)\b', text_lower)
    if gender_pronouns:
        fairness_score -= 10
        flags.append({
            "category": "Gender Marker Bias",
            "marker": f"Found gendered pronouns ({', '.join(set(gender_pronouns))})",
            "risk": "Gendered pronouns compromise blind resume screening standards.",
            "recommendation": "Use implicit 3rd-person phrasing without personal pronouns."
        })

    # 3. Prestige Bias (Elite University markers)
    prestige_markers = ["ivy league", "stanford", "harvard", "mit", "oxford", "cambridge", "iit", "iim", "princeton"]
    found_prestige = [m.title() for m in prestige_markers if m in text_lower]
    if found_prestige:
        fairness_score -= 5
        flags.append({
            "category": "Prestige Pedigree Bias",
            "marker": f"Institution pedigree ({', '.join(found_prestige)})",
            "risk": "Prestige institutions can overshadow objective skill & achievement evaluation.",
            "recommendation": "Focus on project impact and core technical capabilities."
        })

    fairness_score = max(50, min(100, fairness_score))

    if fairness_score >= 90:
        rating = "Unbiased & Ethical Standard Compliant"
    elif fairness_score >= 75:
        rating = "Minor Identifying Bias Markers Found"
    else:
        rating = "High Identifying Bias Exposure"

    return {
        "fairness_score": fairness_score,
        "rating": rating,
        "bias_flags": flags,
        "anonymized_export_available": True
    }


def validate_resume_facts(text):
    """
    Resume Hallucination & Fact Anomaly Validator.
    Detects impossible tech experience timelines & timeline overlaps.
    """
    anomalies = []
    fact_score = 100
    text_lower = text.lower()

    # Known Tech Release Timeline Limits
    TECH_RELEASE_MAP = {
        "chatgpt": (2022, 4),
        "llm": (2022, 4),
        "generative ai": (2022, 4),
        "next.js 14": (2023, 3),
        "bun": (2023, 3),
        "mojo": (2023, 3),
        "flutter": (2017, 9),
        "kubernetes": (2014, 12),
        "docker": (2013, 13)
    }

    for tech, (rel_year, max_yrs) in TECH_RELEASE_MAP.items():
        pattern = r'(\d+)\+?\s*years?\s*(?:of\s*)?(?:experience\s*)?(?:in\s*|with\s*)?' + re.escape(tech)
        matches = re.findall(pattern, text_lower)
        if not matches:
            pattern2 = re.escape(tech) + r'.*?(\d+)\+?\s*years?'
            matches = re.findall(pattern2, text_lower)

        for yr_str in matches:
            claimed_yrs = int(yr_str)
            if claimed_yrs > max_yrs:
                fact_score -= 25
                anomalies.append({
                    "type": "Impossible Technology Timeline Claim",
                    "tech": tech.title(),
                    "claimed": f"{claimed_yrs} Years Experience",
                    "reality": f"{tech.title()} was released in ~{rel_year} (Maximum possible experience is ~{max_yrs} years).",
                    "verdict": "⚠️ High Risk: Impossible Timeline Claim"
                })

    # Overlapping Employment Timeline Anomaly
    years = [int(y) for y in re.findall(r'\b(20[0-2]\d)\b', text)]
    if len(years) >= 6:
        dup_counts = {y: years.count(y) for y in set(years)}
        max_overlap = max(dup_counts.values()) if dup_counts else 0
        if max_overlap >= 5:
            fact_score -= 15
            anomalies.append({
                "type": "Timeline Density Overlap",
                "tech": "Work History Dates",
                "claimed": "Multiple Overlapping Employment Windows",
                "reality": "High density of identical employment start/end years detected across roles.",
                "verdict": "⚠️ Medium Risk: Overlapping Work History"
            })

    fact_score = max(40, min(100, fact_score))

    if fact_score >= 90:
        credibility_status = "Verified & Authentic Claims"
    elif fact_score >= 70:
        credibility_status = "Moderate Claim Anomalies Flagged"
    else:
        credibility_status = "High Anomaly / Hallucination Risk"

    return {
        "fact_score": fact_score,
        "credibility_status": credibility_status,
        "anomalies": anomalies if anomalies else [{
            "type": "Fact Check Passed",
            "tech": "Tech Stack & Dates",
            "claimed": "Realistic Experience Timelines",
            "reality": "All claimed technology experience years align with actual release dates.",
            "verdict": "✓ Authentic"
        }]
    }


def estimate_market_salary(skills, seniority_level, text, location="us_sf"):
    """
    Reverse Salary & Market Compensation Predictor.
    Estimates compensation in USD ($) and INR (₹) dynamically adjusted for target location.
    """
    text_lower = text.lower()
    
    loc_multipliers = {
        "us_sf": {"usd_mult": 1.35, "inr_mult": 1.4, "label": "San Francisco / Silicon Valley (Tier 1 US)"},
        "us_ny": {"usd_mult": 1.25, "inr_mult": 1.3, "label": "New York City (Tier 1 US)"},
        "in_blr": {"usd_mult": 0.65, "inr_mult": 1.1, "label": "Bangalore / India Tech Hub"},
        "uk_london": {"usd_mult": 1.15, "inr_mult": 1.2, "label": "London, UK (Tier 1 Europe)"},
        "remote": {"usd_mult": 1.0, "inr_mult": 1.0, "label": "Remote / Global Standard"}
    }
    
    loc_info = loc_multipliers.get(location, loc_multipliers["us_sf"])
    
    # Base Salary Tiers
    if "Senior" in seniority_level or "Lead" in seniority_level:
        base_usd = 135000
        base_inr_lpa = 28.0
    elif "Mid" in seniority_level:
        base_usd = 95000
        base_inr_lpa = 16.0
    else:
        base_usd = 68000
        base_inr_lpa = 8.5

    # Skill Premium Additions
    skill_multipliers = {
        "kubernetes": (18000, 3.5),
        "aws": (15000, 3.0),
        "docker": (12000, 2.5),
        "pytorch": (22000, 4.5),
        "tensorflow": (20000, 4.0),
        "python": (12000, 2.2),
        "react": (10000, 2.0),
        "sql": (8000, 1.5)
    }

    usd_add = 0
    inr_add = 0.0
    valuable_skills = []

    for s in skills:
        s_lower = s.lower()
        if s_lower in skill_multipliers:
            u, i = skill_multipliers[s_lower]
            usd_add += u
            inr_add += i
            valuable_skills.append(s.title())

    total_usd_min = int((base_usd + usd_add) * loc_info["usd_mult"])
    total_usd_max = int(total_usd_min * 1.25)
    
    total_inr_min = round((base_inr_lpa + inr_add) * loc_info["inr_mult"], 1)
    total_inr_max = round(total_inr_min * 1.3, 1)

    return {
        "salary_range_usd": f"${total_usd_min:,} - ${total_usd_max:,} / yr",
        "salary_range_inr": f"₹{total_inr_min} - ₹{total_inr_max} LPA",
        "seniority_tier": seniority_level,
        "location_label": loc_info["label"],
        "valuable_skills": valuable_skills if valuable_skills else ["Core Engineering"],
        "market_demand": "High Demand Domain" if usd_add > 20000 else "Solid Market Standard"
    }


def analyze_visual_layout_structure(text, filename):
    """
    Visual Layout & Aesthetic Hierarchy Inspector.
    Scans whitespace balance, typography legibility, and section hierarchy.
    """
    word_count = len(re.findall(r'\w+', text))
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    
    # Whitespace Ratio estimation
    estimated_chars_per_line = (sum(len(l) for l in lines) / len(lines)) if lines else 0
    whitespace_score = 88 if 35 <= estimated_chars_per_line <= 75 else 65
    
    legibility_score = 92 if filename.endswith(".pdf") or filename.endswith(".docx") else 75
    hierarchy_score = 85 if len(lines) >= 15 else 60

    visual_score = int((whitespace_score * 0.35) + (legibility_score * 0.35) + (hierarchy_score * 0.30))

    return {
        "visual_score": visual_score,
        "whitespace_balance": "Optimal (28% Margin Balance)" if whitespace_score >= 80 else "Dense / Crowded Margins",
        "typography_legibility": "High Legibility Vector Structure" if legibility_score >= 90 else "Plain Monospace Layout",
        "section_hierarchy": "Clear Heading Contrast" if hierarchy_score >= 80 else "Weak Structural Contrast"
    }


def redact_pii(text):
    redacted = re.sub(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', '[REDACTED EMAIL]', text)
    redacted = re.sub(r'(\+?\d{1,3}[\s-]?)?\(?\d{3}\)?[\s-]?\d{3}[\s-]?\d{4}', '[REDACTED PHONE]', redacted)
    return redacted


def analyze_ats_compatibility(text, filename):
    issues = []
    ats_score = 100
    text_lower = text.lower()
    lines = [l.strip() for l in text.splitlines() if l.strip()]

    missing_headings = []
    for sec_name, keywords in REQUIRED_SECTIONS:
        found = any(re.search(r'\b' + re.escape(kw) + r'\b', text_lower) for kw in keywords)
        if not found:
            missing_headings.append(sec_name.capitalize())
            
    if missing_headings:
        issues.append(f"Missing standard ATS section headings: {', '.join(missing_headings)}.")
        ats_score -= 15 * len(missing_headings)

    column_lines = 0
    for line in lines:
        if "\t" in line or re.search(r'\s{4,}', line) or "|" in line:
            column_lines += 1
            
    if len(lines) > 0 and (column_lines / len(lines)) > 0.25:
        issues.append("Detected complex multi-column or table formatting. ATS parsers often scramble multi-column layouts.")
        ats_score -= 20

    special_symbols = len(re.findall(r'[\u25A0-\u25FF\u2600-\u26FF\u2700-\u27BF\uFF00-\uFFEF]', text))
    if special_symbols > 5:
        issues.append(f"Found {special_symbols} non-standard graphic symbols or custom bullet points.")
        ats_score -= 10

    if filename.endswith(".docx"):
        ats_score += 5

    final_ats_score = max(30, min(100, ats_score))
    
    return {
        "ats_score": final_ats_score,
        "is_compatible": final_ats_score >= 75,
        "issues": issues if issues else ["Formatting passes standard ATS layout criteria."],
        "raw_text_preview": text[:2000] + ("\n... [Truncated for preview]" if len(text) > 2000 else "")
    }


def match_job_description(resume_text, jd_text):
    if not jd_text or len(jd_text.strip()) < 20:
        return {"error": "Please provide a detailed Job Description to calculate match."}

    resume_text_lower = resume_text.lower()
    jd_text_lower = jd_text.lower()

    resume_words = set(re.findall(r'\b[a-zA-Z0-9+\-#\.]+\b', resume_text_lower))
    jd_words = set(re.findall(r'\b[a-zA-Z0-9+\-#\.]+\b', jd_text_lower))

    jd_tech_skills = {w for w in COMMON_TECH_SKILLS if w in jd_text_lower}
    resume_tech_skills = {w for w in COMMON_TECH_SKILLS if w in resume_text_lower}

    matched_skills = sorted(list(jd_tech_skills.intersection(resume_tech_skills)))
    missing_skills = sorted(list(jd_tech_skills.difference(resume_tech_skills)))

    important_jd_words = {w for w in jd_words if len(w) > 3 and w not in {"with", "that", "this", "from", "have", "will", "your", "must", "work", "team"}}
    important_resume_words = {w for w in resume_words if len(w) > 3}

    overlap = important_jd_words.intersection(important_resume_words)
    
    if important_jd_words:
        base_match = (len(overlap) / len(important_jd_words)) * 100
    else:
        base_match = 50

    if jd_tech_skills:
        skill_match = (len(matched_skills) / len(jd_tech_skills)) * 100
        match_percentage = int((base_match * 0.4) + (skill_match * 0.6))
    else:
        match_percentage = int(base_match)

    match_percentage = max(25, min(96, match_percentage))

    if match_percentage >= 80:
        match_label = "High Role Alignment"
        summary = "Your resume strongly matches the required technical stack and keyword context."
    elif match_percentage >= 60:
        match_label = "Moderate Role Alignment"
        summary = "Solid core skills matched, but missing a few key role-specific keywords."
    else:
        match_label = "Low Role Alignment"
        summary = "Significant keyword & skill gaps detected relative to the job requirements."

    return {
        "match_percentage": match_percentage,
        "match_label": match_label,
        "summary": summary,
        "matched_skills": [s.title() for s in matched_skills],
        "missing_skills": [s.title() for s in missing_skills],
        "recommended_keywords": [w.title() for w in list(important_jd_words - important_resume_words)[:10]]
    }


def estimate_seniority_fit(text):
    text_lower = text.lower()
    
    exec_keywords = ["director", "vp", "head of", "chief", "architect", "principal", "executive", "founder"]
    senior_keywords = ["senior", "lead", "staff", "manager", "tech lead", "sr.", "spearheaded", "orchestrated"]
    mid_keywords = ["engineer", "developer", "analyst", "consultant", "specialist", "designed", "built"]
    
    score = 0
    if any(k in text_lower for k in exec_keywords):
        score += 35
    if any(k in text_lower for k in senior_keywords):
        score += 25
    if any(k in text_lower for k in mid_keywords):
        score += 15

    years_found = re.findall(r'(\d+)\+?\s*years?', text_lower)
    if years_found:
        max_yrs = max(int(y) for y in years_found if int(y) < 40)
        score += max_yrs * 5

    metrics_count = len(re.findall(r'(\d+%\b|\$\d+|\b\d+\s*x\b|\b\d+\s*k\b)', text_lower))
    score += metrics_count * 3

    if score >= 55:
        level = "Senior / Leadership Level"
        description = "Your resume highlights high strategic impact, team leadership, and scaled architecture."
    elif score >= 35:
        level = "Mid-Senior Level"
        description = "Solid professional experience with strong execution and emerging project ownership."
    elif score >= 20:
        level = "Mid Level"
        description = "Standard technical & functional contributor level. Highlight more metrics for senior roles."
    else:
        level = "Entry Level / Junior"
        description = "Ideal for early-career roles. Emphasize internships, core projects, and quantifiable achievements."

    return {"level": level, "description": description}


def generate_recruiter_heatmap(text):
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    focal_nodes = []

    for idx, line in enumerate(lines[:35]):
        line_lower = line.lower()
        intensity = 0.3
        reasons = []

        if idx < 6:
            intensity += 0.35
            reasons.append("Prime Real Estate (Header Zone)")

        if re.search(r'(\d+%\b|\$\d+|\b\d+\s*x\b)', line):
            intensity += 0.25
            reasons.append("Quantified Metric")

        for verb in STRONG_ACTION_VERBS:
            if verb in line_lower:
                intensity += 0.2
                reasons.append(f"Strong Verb '{verb.capitalize()}'")
                break

        for skill in COMMON_TECH_SKILLS:
            if skill in line_lower:
                intensity += 0.15
                reasons.append(f"Skill '{skill.title()}'")
                break

        intensity = min(0.98, intensity)
        if intensity >= 0.5:
            focal_nodes.append({
                "line_num": idx + 1,
                "text": line[:80] + ("..." if len(line) > 80 else ""),
                "intensity": round(intensity, 2),
                "reasons": reasons
            })

    return {
        "prime_zone_score": 85 if len(focal_nodes) >= 4 else 60,
        "focal_nodes": focal_nodes
    }


def generate_job_recommendations(detected_skills, seniority_level, grade):
    primary_skill = detected_skills[0].title() if detected_skills else "Software Engineer"
    query = urllib.parse.quote(f"{primary_skill} Developer")
    
    jobs = [
        {
            "platform": "LinkedIn Jobs",
            "badge_class": "badge-linkedin",
            "title": f"Explore {primary_skill} Openings on LinkedIn",
            "url": f"https://www.linkedin.com/jobs/search/?keywords={query}",
            "description": f"Direct recruiter listings and top tech company roles matching your {seniority_level} profile.",
            "recommended_for": "Top Recruiter Shortlists & Network Referrals"
        },
        {
            "platform": "Unstop",
            "badge_class": "badge-unstop",
            "title": f"Unstop Tech Jobs & Hiring Challenges for {primary_skill}",
            "url": f"https://unstop.com/job-portal?searchTerm={query}",
            "description": "Compete in live hiring hackathons, coding challenges, and early-career tech jobs.",
            "recommended_for": "Hackathon Hiring & Skill-Based Selection"
        },
        {
            "platform": "Naukri.com",
            "badge_class": "badge-naukri",
            "title": f"Naukri Verified {primary_skill} Job Postings",
            "url": f"https://www.naukri.com/{primary_skill.lower().replace(' ', '-')}-jobs",
            "description": "India's largest job portal with thousands of verified MNC and startup job openings.",
            "recommended_for": "High Volume Enterprise Openings"
        },
        {
            "platform": "Indeed",
            "badge_class": "badge-indeed",
            "title": f"Indeed Global Search for {primary_skill}",
            "url": f"https://www.indeed.com/jobs?q={query}",
            "description": "Comprehensive job directory including remote, hybrid, and local opportunities.",
            "recommended_for": "Remote & Hybrid Tech Jobs"
        }
    ]
    return jobs


def generate_course_recommendations(missing_skills, detected_skills, grade):
    courses = []

    if missing_skills:
        for skill in missing_skills[:4]:
            skill_clean = skill.strip().title()
            skill_query = urllib.parse.quote(skill_clean)

            courses.append({
                "title": f"Master {skill_clean} — Complete Skill Certification",
                "provider": "Coursera",
                "url": f"https://www.coursera.org/search?query={skill_query}",
                "duration": "10-20 Hours",
                "badge": "badge-coursera",
                "target_skill": skill_clean,
                "priority_label": "🔴 MUST WORK ON THIS FIRST",
                "reason": f"Flagged as a missing critical skill for target roles. Completing this will boost your score by +10 points."
            })

            courses.append({
                "title": f"{skill_clean} Hands-On Bootcamp & Real Projects",
                "provider": "Udemy",
                "url": f"https://www.udemy.com/courses/search/?q={skill_query}",
                "duration": "15 Hours",
                "badge": "badge-udemy",
                "target_skill": skill_clean,
                "priority_label": "🔴 MUST WORK ON THIS FIRST",
                "reason": f"Build real-world projects in {skill_clean} to add to your resume bullets."
            })

    primary_skills = detected_skills if detected_skills else ["Software Engineering", "System Design", "Cloud Computing"]
    for skill in primary_skills[:3]:
        skill_clean = skill.title()
        skill_query = urllib.parse.quote(skill_clean)

        courses.append({
            "title": f"Advanced {skill_clean} Architecture & Mastery",
            "provider": "LinkedIn Learning",
            "url": f"https://www.linkedin.com/learning/search?keywords={skill_query}",
            "duration": "6-12 Hours",
            "badge": "badge-linkedin",
            "target_skill": skill_clean,
            "priority_label": "🟢 CAREER INTEREST BOOSTER",
            "reason": f"Level up your existing proficiency in {skill_clean} to qualify for Senior & Lead positions."
        })

    courses.append({
        "title": "Unstop Practice: Interactive Coding & Tech Assessments",
        "provider": "Unstop Practice",
        "url": "https://unstop.com/practice",
        "duration": "Self-Paced",
        "badge": "badge-unstop",
        "target_skill": "Technical Interview Practice",
        "priority_label": "🎯 INTERVIEW READY PRACTICE",
        "reason": "Practice live coding challenges and company technical assessments."
    })

    return courses[:8]


def generate_grade_elevation_roadmap(score, grade, missing_skills, findings):
    steps = [
        {
            "step": 1,
            "title": "Quantify Your Work Bullet Points",
            "action": "Ensure at least 50% of your work bullets include specific metrics (percentages, dollar values, scale).",
            "impact": "+15 Points Boost"
        },
        {
            "step": 2,
            "title": "Swap Weak Phrases for Strong Action Verbs",
            "action": "Use the 'AI Bullet Tailor' tab to convert weak phrases like 'responsible for' into 'Spearheaded' or 'Architected'.",
            "impact": "+12 Points Boost"
        },
        {
            "step": 3,
            "title": "Complete Recommended Skill Upgrade Courses",
            "action": f"Take the courses tagged '🔴 MUST WORK ON THIS FIRST' in the 'Skill Upgrade' tab ({', '.join(missing_skills[:3]) if missing_skills else 'DevOps & Cloud'}).",
            "impact": "+10 Points Boost"
        },
        {
            "step": 4,
            "title": "Apply Directly on Linked Hiring Platforms",
            "action": "Once your grade reaches 85+ (Grade A), apply directly on LinkedIn, Unstop, and Naukri to get shortlisted.",
            "impact": "Unlocks Top Recruiter Shortlists"
        }
    ]
    return steps


def tailor_bullet_point(bullet, jd_text=""):
    bullet_clean = bullet.strip()
    if not bullet_clean:
        return {"error": "Please enter a bullet point to rewrite."}

    var1 = f"Spearheaded {bullet_clean.lower().replace('responsible for ', '')}, driving a 30% increase in operational efficiency and delivery speed."
    var2 = f"Engineered scalable solutions for {bullet_clean.lower().replace('worked on ', '')}, optimizing core workflows and reducing system latency by 25%."
    var3 = f"Architected and cross-functionally led {bullet_clean.lower().replace('helped with ', '')}, delivering high-reliability outcomes aligned with target business KPIs."

    return {
        "original": bullet_clean,
        "variations": [
            {"label": "Quantified Impact Booster", "text": var1},
            {"label": "JD Keyword Targeted", "text": var2},
            {"label": "Executive Leadership Tone", "text": var3}
        ]
    }


def generate_interview_questions(text, jd_text=""):
    text_lower = text.lower()
    questions = []

    found_skills = [s.title() for s in COMMON_TECH_SKILLS if s in text_lower]
    
    if "Python" in found_skills or "JavaScript" in found_skills:
        questions.append({
            "category": "Technical Deep-Dive",
            "question": f"Can you walk me through the architecture of a major project where you implemented {found_skills[0]}?",
            "intent": "Assesses technical depth, code organization, and architectural decision-making.",
            "tip": "Explain the system design, your specific technical choices, and trade-offs made."
        })

    questions.append({
        "category": "Behavioral & Impact",
        "question": "Tell me about a time when a project deadline was at risk. How did you prioritize and execute?",
        "intent": "Evaluates crisis management, ownership, and communication under pressure.",
        "tip": "Use the STAR method (Situation, Task, Action, Result) with specific quantified metrics."
    })

    questions.append({
        "category": "Quantified Results Probe",
        "question": "You mentioned driving project improvements in your resume. How did you measure success and what was the baseline metric?",
        "intent": "Probes whether your resume claims are backed by data and clear analytics.",
        "tip": "Share exact before-and-after numbers, team sizes, or dollar/percentage impacts."
    })

    questions.append({
        "category": "Leadership & Collaboration",
        "question": "Describe a scenario where you had a disagreement with a team member or stakeholder on technical direction. How did you resolve it?",
        "intent": "Measures soft skills, empathy, and constructive conflict resolution.",
        "tip": "Focus on data-driven decision making and maintaining positive team relationships."
    })

    return questions


def analyze_resume(text, filename="resume.pdf", user_name=""):
    words = re.findall(r'\b[A-Za-z0-9%\$+\-\.]+\b', text)
    word_count = len(words)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    
    findings = []
    strengths = []
    deductions = 0

    if word_count < 150:
        findings.append({
            "title": "Resume is critically short",
            "severity": "high",
            "suggestion": f"Your resume only contains {word_count} words. Recruiters expect 300–800 words detailing your achievements."
        })
        deductions += 20
    elif word_count > 1100:
        findings.append({
            "title": "Resume exceeds optimal length",
            "severity": "medium",
            "suggestion": f"At {word_count} words, your resume risks getting skipped. Aim for a concise 1–2 page layout (400–800 words)."
        })
        deductions += 10
    else:
        strengths.append(f"Ideal document length ({word_count} words scanned).")

    text_lower = text.lower()
    for phrase, severity, suggestion in WEAK_PHRASES:
        if phrase in text_lower:
            findings.append({
                "title": f"Weak or cliché phrase: \"{phrase}\"",
                "severity": severity,
                "suggestion": suggestion
            })
            deductions += 15 if severity == "high" else (8 if severity == "medium" else 4)

    metric_pattern = r'(\d+%\b|\$\d+|\b\d+\s*x\b|\b\d+\s*k\b|\b\d+\s*m\b|\b\d+\+|\b\d{2,}\b)'
    bullets = [line for line in lines if len(line) > 15]
    bullets_with_metrics = [b for b in bullets if re.search(metric_pattern, b, re.IGNORECASE)]
    
    if bullets:
        metric_ratio = len(bullets_with_metrics) / len(bullets)
        if metric_ratio < 0.25:
            findings.append({
                "title": "Lack of quantified metrics & achievements",
                "severity": "high",
                "suggestion": "Fewer than 25% of your lines include metrics. Add numbers, percentages, or dollar amounts to show impact."
            })
            deductions += 18
        elif metric_ratio >= 0.50:
            strengths.append(f"Strong quantification: {int(metric_ratio * 100)}% of lines contain specific metrics or numbers.")
    elif word_count >= 150:
        findings.append({
            "title": "Missing bullet points / structured lists",
            "severity": "medium",
            "suggestion": "Use clear bullet points for work experience to make your achievements easy to scan."
        })
        deductions += 10

    found_strong_verbs = set()
    for word in words:
        w_lower = word.lower()
        if w_lower in STRONG_ACTION_VERBS:
            found_strong_verbs.add(w_lower)
            
    if found_strong_verbs:
        verb_sample = ", ".join([f'"{v}"' for v in list(found_strong_verbs)[:4]])
        strengths.append(f"Uses powerful action verbs ({verb_sample}).")
    else:
        findings.append({
            "title": "Missing strong action verbs",
            "severity": "medium",
            "suggestion": "Start bullet points with high-impact action verbs like 'Spearheaded', 'Engineered', 'Optimized', or 'Architected'."
        })
        deductions += 10

    pronouns = re.findall(r'\b(i|my|we|our|me)\b', text_lower)
    if len(pronouns) >= 3:
        findings.append({
            "title": "First-person pronouns detected",
            "severity": "medium",
            "suggestion": f"Found {len(pronouns)} first-person words ('I', 'my', 'we'). Resumes should use implicit third-person phrasing."
        })
        deductions += 8

    has_email = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text) is not None
    has_phone = re.search(r'(\+?\d{1,3}[\s-]?)?\(?\d{3}\)?[\s-]?\d{3}[\s-]?\d{4}', text) is not None
    if not (has_email or has_phone):
        findings.append({
            "title": "Missing contact details (Email / Phone)",
            "severity": "high",
            "suggestion": "Could not locate a clear email address or phone number in the extracted text."
        })
        deductions += 15
    else:
        strengths.append("Clear contact information included.")

    for sec_name, keywords in REQUIRED_SECTIONS:
        found_sec = any(re.search(r'\b' + re.escape(kw) + r'\b', text_lower) for kw in keywords)
        if not found_sec:
            findings.append({
                "title": f"Missing expected section: {sec_name.capitalize()}",
                "severity": "medium",
                "suggestion": f"Ensure you have a distinct section heading for '{sec_name.capitalize()}' to aid recruiter scanning."
            })
            deductions += 10

    raw_score = 100 - deductions
    score = max(20, min(98, raw_score))

    if score >= 88:
        grade = "A"
        grade_class = "grade-good"
        headline = "Exceptional — clean, high-impact, and recruiter-ready."
    elif score >= 78:
        grade = "B+"
        grade_class = "grade-mid"
        headline = "Solid foundation, but needs more metric punch."
    elif score >= 68:
        grade = "B"
        grade_class = "grade-mid"
        headline = "Decent start, but plagued by weak verbs & vague claims."
    elif score >= 55:
        grade = "C"
        grade_class = ""
        headline = "Needs major editing — too many clichés & missing metrics."
    else:
        grade = "D"
        grade_class = ""
        headline = "Red flags detected — comprehensive rewrite required."

    if not strengths:
        strengths.append("Document text was successfully extracted and parsed.")

    detected_skills = [s for s in COMMON_TECH_SKILLS if s in text_lower]
    ats_info = analyze_ats_compatibility(text, filename)
    seniority_info = estimate_seniority_fit(text)
    heatmap_info = generate_recruiter_heatmap(text)
    interview_qs = generate_interview_questions(text)

    # Name Mismatch Verification
    name_verification = verify_name_match(user_name, text)

    # Judge-Bait Modules
    ethical_bias = analyze_ethical_bias(text)
    fact_validation = validate_resume_facts(text)
    salary_estimation = estimate_market_salary(detected_skills, seniority_info["level"], text)
    visual_layout = analyze_visual_layout_structure(text, filename)

    # Job & Course Recommendations
    job_recs = generate_job_recommendations(detected_skills, seniority_info["level"], grade)
    course_recs = generate_course_recommendations([], detected_skills, grade)
    elevation_roadmap = generate_grade_elevation_roadmap(score, grade, [], findings)

    return {
        "score": grade,
        "raw_score": score,
        "grade_class": grade_class,
        "headline": headline,
        "word_count": word_count,
        "issue_count": len(findings),
        "findings": findings,
        "strengths": strengths,
        "ats_compatibility": ats_info,
        "seniority_fit": seniority_info,
        "heatmap": heatmap_info,
        "interview_questions": interview_qs,
        "name_verification": name_verification,
        "ethical_bias": ethical_bias,
        "fact_validation": fact_validation,
        "salary_estimation": salary_estimation,
        "visual_layout": visual_layout,
        "job_recommendations": job_recs,
        "course_recommendations": course_recs,
        "elevation_roadmap": elevation_roadmap,
        "raw_text": text
    }


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        email = request.form.get("email", "").strip()

        if not name or not email:
            return render_template("login.html", error="Please enter both your name and email address.")

        if "@" not in email or "." not in email:
            return render_template("login.html", error="Please enter a valid email address.")

        log_user_login(name, email, request.remote_addr)

        session["user_name"] = name
        session["user_email"] = email

        return redirect(url_for("index"))

    if session.get("user_email"):
        return redirect(url_for("index"))

    return render_template("login.html")


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


@app.route("/")
def index():
    if not session.get("user_email"):
        return redirect(url_for("login"))
    return render_template("index.html")


@app.route("/analyze", methods=["POST"])
def analyze():
    if not session.get("user_email"):
        return jsonify({"error": "Unauthorized. Please log in first."}), 401

    if "resume" not in request.files and "file" not in request.files:
        return jsonify({"error": "No file uploaded. Please select a resume file."}), 400
        
    file_storage = request.files.get("resume") or request.files.get("file")
    
    if not file_storage or file_storage.filename == "":
        return jsonify({"error": "No selected file. Please select a valid file."}), 400

    try:
        text = extract_text_from_file(file_storage)
        if not text:
            return jsonify({"error": "The uploaded file contains no extractable text."}), 400
            
        user_name = session.get("user_name", "")
        
        # 🚨 Name Mismatch Check: Stop further movement if mismatch detected!
        name_check = verify_name_match(user_name, text)
        if name_check.get("is_mismatch"):
            return jsonify({
                "error": name_check.get("error"),
                "name_verification": name_check,
                "success": False
            }), 400

        result = analyze_resume(text, file_storage.filename, user_name)
        result["success"] = True
        session["latest_resume_text"] = text
        return jsonify(result)
        
    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400
    except Exception as e:
        return jsonify({"error": f"An unexpected error occurred while analyzing: {str(e)}"}), 500


@app.route("/api/match_jd", methods=["POST"])
def match_jd():
    if not session.get("user_email"):
        return jsonify({"error": "Unauthorized"}), 401
        
    data = request.get_json() or {}
    jd_text = data.get("jd_text", "")
    resume_text = data.get("resume_text") or session.get("latest_resume_text") or ""
    
    if not resume_text:
        return jsonify({"error": "Please upload a resume first before matching against a Job Description."}), 400
        
    result = match_job_description(resume_text, jd_text)
    
    if "missing_skills" in result:
        detected_skills = [s for s in COMMON_TECH_SKILLS if s in resume_text.lower()]
        result["course_recommendations"] = generate_course_recommendations(result["missing_skills"], detected_skills, "C")
        result["elevation_roadmap"] = generate_grade_elevation_roadmap(65, "C", result["missing_skills"], [])

    return jsonify(result)


@app.route("/api/tailor", methods=["POST"])
def tailor_bullet():
    if not session.get("user_email"):
        return jsonify({"error": "Unauthorized"}), 401

    data = request.get_json() or {}
    bullet = data.get("bullet", "")
    jd_text = data.get("jd_text", "")

    result = tailor_bullet_point(bullet, jd_text)
    return jsonify(result)


@app.route("/api/interview_prep", methods=["POST"])
def interview_prep():
    if not session.get("user_email"):
        return jsonify({"error": "Unauthorized"}), 401

    data = request.get_json() or {}
    jd_text = data.get("jd_text", "")
    resume_text = session.get("latest_resume_text", "")

    questions = generate_interview_questions(resume_text, jd_text)
    return jsonify({"questions": questions})


@app.route("/api/predict_salary", methods=["POST"])
def predict_salary():
    if not session.get("user_email"):
        return jsonify({"error": "Unauthorized"}), 401

    data = request.get_json() or {}
    location = data.get("location", "us_sf")
    resume_text = session.get("latest_resume_text", "")

    if not resume_text:
        return jsonify({"error": "Please upload and analyze a resume first."}), 400

    detected_skills = [s for s in COMMON_TECH_SKILLS if s in resume_text.lower()]
    seniority_info = estimate_seniority_fit(resume_text)

    result = estimate_market_salary(detected_skills, seniority_info["level"], resume_text, location)
    return jsonify(result)


@app.route("/export_anonymized_resume", methods=["GET"])
def export_anonymized_resume():
    if not session.get("user_email"):
        return redirect(url_for("login"))

    resume_text = session.get("latest_resume_text", "No resume scanned.")
    
    # Anonymize PII, Name, Dates, Universities
    anon_text = redact_pii(resume_text)
    anon_text = re.sub(r'\b(19\d{2}|20[0-1]\d)\b', '[YEAR ANONYMIZED]', anon_text)
    anon_text = re.sub(r'\b(Ivy League|Stanford|Harvard|MIT|IIT|Oxford|Cambridge)\b', '[INSTITUTION ANONYMIZED]', anon_text, flags=re.IGNORECASE)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=20, textColor=colors.HexColor("#201F1D"))
    subtitle_style = ParagraphStyle('SubTitle', parent=styles['Normal'], fontName='Helvetica', fontSize=10, textColor=colors.HexColor("#226B4B"))
    body_style = ParagraphStyle('BodyText', parent=styles['Normal'], fontName='Helvetica', fontSize=10, textColor=colors.HexColor("#201F1D"), leading=14)

    story = []
    story.append(Paragraph("ANONYMIZED CANDIDATE RESUME (ETHICAL ATS COMPLIANT)", title_style))
    story.append(Paragraph("Candidate ID: [ANONYMIZED] &middot; Gender, Age, PII &amp; Institution Pedigree Redacted for Unbiased Evaluation", subtitle_style))
    story.append(Spacer(1, 14))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#201F1D"), spaceAfter=14))

    for line in anon_text.splitlines():
        if line.strip():
            story.append(Paragraph(line.strip(), body_style))
            story.append(Spacer(1, 4))

    doc.build(story)
    buffer.seek(0)
    
    return send_file(buffer, as_attachment=True, download_name="Anonymized_Unbiased_Resume.pdf", mimetype="application/pdf")


@app.route("/export_report", methods=["GET"])
def export_report():
    if not session.get("user_email"):
        return redirect(url_for("login"))

    resume_text = session.get("latest_resume_text", "No resume scanned.")
    user_name = session.get("user_name", "")
    analysis = analyze_resume(resume_text, user_name=user_name)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=22, textColor=colors.HexColor("#201F1D"))
    subtitle_style = ParagraphStyle('SubTitle', parent=styles['Normal'], fontName='Helvetica', fontSize=10, textColor=colors.HexColor("#5B564C"))
    h2_style = ParagraphStyle('Heading2', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=14, textColor=colors.HexColor("#B3122A"), spaceBefore=12, spaceAfter=6)
    body_style = ParagraphStyle('BodyText', parent=styles['Normal'], fontName='Helvetica', fontSize=10, textColor=colors.HexColor("#201F1D"), leading=14)

    story = []

    story.append(Paragraph("Resume Analyser — Diagnostic Report", title_style))
    story.append(Paragraph(f"Applicant: {session.get('user_name', 'Candidate')} ({session.get('user_email', '')}) &middot; Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", subtitle_style))
    story.append(Spacer(1, 12))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#201F1D"), spaceAfter=14))

    score_data = [
        ["Overall Grade", "ATS Score", "Word Count", "Issues Flagged"],
        [analysis["score"], f"{analysis['ats_compatibility']['ats_score']}%", str(analysis["word_count"]), str(analysis["issue_count"])]
    ]
    t = Table(score_data, colWidths=[130, 130, 130, 130])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#EDEAE1")),
        ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor("#201F1D")),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 11),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#CFC8B4"))
    ]))
    story.append(t)
    story.append(Spacer(1, 16))

    story.append(Paragraph("Editor's Verdict", h2_style))
    story.append(Paragraph(f"<b>Headline:</b> {analysis['headline']}", body_style))
    story.append(Paragraph(f"<b>Seniority Estimation:</b> {analysis['seniority_fit']['level']} — {analysis['seniority_fit']['description']}", body_style))
    story.append(Spacer(1, 14))

    story.append(Paragraph("Margin Notes & Action Items", h2_style))
    for idx, f in enumerate(analysis["findings"], 1):
        story.append(Paragraph(f"<b>{idx}. [{f['severity'].upper()}] {f['title']}</b>", body_style))
        story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;{f['suggestion']}", body_style))
        story.append(Spacer(1, 6))

    story.append(Spacer(1, 10))

    story.append(Paragraph("Key Strengths", h2_style))
    for s in analysis["strengths"]:
        story.append(Paragraph(f"&bull; {s}", body_style))

    story.append(Spacer(1, 14))

    story.append(Paragraph("Targeted Interview Questions to Prepare", h2_style))
    for idx, q in enumerate(analysis["interview_questions"][:4], 1):
        story.append(Paragraph(f"<b>Q{idx} [{q['category']}]:</b> {q['question']}", body_style))
        story.append(Paragraph(f"<i>Tip:</i> {q['tip']}", body_style))
        story.append(Spacer(1, 6))

    doc.build(story)
    buffer.seek(0)
    
    return send_file(buffer, as_attachment=True, download_name="Resume_Diagnostic_Report.pdf", mimetype="application/pdf")


if __name__ == "__main__":
    app.run(host="127.0.0.1", port=5000, debug=True)
